import streamlit as st
import uuid
from datetime import datetime
from typing import List

# Import các thư viện xử lý tài liệu và kết nối CSDL
# Giả định các import này hoạt động chính xác trong môi trường của bạn
try:
    import PyPDF2
    import docx
    from src.database.pinecone_client import upsert_chunk_texts
    from src.database.mongo_client import document_collection
    from src.models.document import Document
    from src.configs.settings import app_config, pinecone_config
except ImportError as e:
    st.error(f"Lỗi import: {e}")
    st.warning("Vui lòng đảm bảo rằng cấu trúc file và các thư viện cần thiết đã được cài đặt.")
    st.stop()

# --- Cấu hình trang ---
st.set_page_config(
    page_title="Tải tài liệu - THCS Nhân Chính",
    page_icon="📚",
    layout="centered"
)

# --- Các hàm xử lý (Giữ nguyên logic gốc) ---

def extract_text_from_pdf(file) -> str:
    """Trích xuất văn bản từ file PDF."""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file) -> str:
    """Trích xuất văn bản từ file Word."""
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_txt(file) -> str:
    """Trích xuất văn bản từ file text."""
    content = file.read()
    # Thử decode bằng utf-8, nếu lỗi thì thử latin-1
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        file.seek(0)
        return content.decode('latin-1', errors='ignore')

def extract_text(file, file_type: str) -> str:
    """Trích xuất văn bản dựa trên loại file."""
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.txt': extract_text_from_txt,
        '.md': extract_text_from_txt,
    }
    if file_type in extractors:
        return extractors[file_type](file)
    raise ValueError(f"Loại file không được hỗ trợ: {file_type}")

def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    """Chia văn bản thành các đoạn nhỏ có chồng lấn."""
    if not text or not text.strip():
        return []
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start += (chunk_size - overlap)
        
    return [chunk.strip() for chunk in chunks if chunk.strip()]

def process_document(file, filename: str, topic: str) -> tuple:
    """Xử lý tài liệu, trả về các đoạn nhỏ và metadata."""
    file_type = '.' + filename.split('.')[-1].lower()
    
    text = extract_text(file, file_type)
    if not text.strip():
        raise ValueError("Không thể trích xuất nội dung từ file hoặc file trống.")
        
    chunks = split_text_into_chunks(text)
    if not chunks:
        raise ValueError("Nội dung file quá ngắn để có thể chia nhỏ.")

    document = Document(
        document_id=str(uuid.uuid4()),
        name=filename,
        topic=topic,
        file_type=file_type,
        file_size=file.tell(),
        chunk_count=len(chunks)
    )
    return chunks, document

@st.cache_data(ttl=60) # Cache kết quả trong 60 giây
def get_documents_cached():
    """Lấy danh sách tài liệu từ CSDL với cache."""
    try:
        # Giả định hàm này trả về một list các dictionary
        return document_collection.get_all_documents()
    except Exception as e:
        st.error(f"Không thể tải danh sách tài liệu: {e}")
        return []

# --- Giao diện người dùng ---

st.title("📚 Tải Lên Tài Liệu")
st.caption("Thêm kiến thức mới cho hệ thống AI từ các file tài liệu.")

# 1. Chọn chủ đề
st.subheader("1. Chọn chủ đề cho tài liệu")
topic_options = {
    pinecone_config.NAME_SPACE.THONG_TIN_TRUONG.value: "🏫 Thông tin trường"
}
selected_topic = st.selectbox(
    "Chủ đề:",
    options=list(topic_options.keys()),
    format_func=lambda x: topic_options.get(x, "Không xác định"),
    label_visibility="collapsed"
)

# 2. Tải file
st.subheader("2. Tải lên tài liệu của bạn")
uploaded_files = st.file_uploader(
    "Hỗ trợ PDF, DOCX, TXT, MD. Có thể chọn nhiều file.",
    type=['pdf', 'docx', 'txt', 'md'],
    accept_multiple_files=True
)

# 3. Xử lý
if uploaded_files:
    st.subheader("3. Bắt đầu xử lý")
    
    # Hiển thị các file đã chọn trong một expander
    with st.expander(f"Xem {len(uploaded_files)} file đã chọn"):
        for file in uploaded_files:
            st.info(f"📄 **{file.name}** ({file.size / 1024:.1f} KB)")

    if st.button("🚀 Xử lý các tài liệu đã tải lên", type="primary", use_container_width=True):
        progress_bar = st.progress(0, text="Bắt đầu quá trình xử lý...")
        success_count = 0
        
        for i, file in enumerate(uploaded_files):
            # Cập nhật thanh tiến trình
            progress_text = f"Đang xử lý file {i+1}/{len(uploaded_files)}: {file.name}"
            progress_bar.progress((i + 1) / len(uploaded_files), text=progress_text)
            
            try:
                # Đảm bảo con trỏ file ở đầu
                file.seek(0)
                
                # Xử lý file
                chunks, document = process_document(file, file.name, selected_topic)
                
                # Lưu vào CSDL
                # upsert_chunk_texts(chunks, selected_topic, document)
                # document_collection.create_document(
                #     document_id=document.document_id,
                #     name=document.name,
                #     topic=document.topic,
                #     file_type=document.file_type,
                #     file_size=document.file_size,
                #     chunk_count=document.chunk_count
                # )
                
                st.success(f"✅ **{file.name}**: Xử lý thành công ({document.chunk_count} phần).")
                success_count += 1
            
            except Exception as e:
                st.error(f"❌ **{file.name}**: Đã xảy ra lỗi - {e}")

        # Thông báo kết quả cuối cùng
        progress_bar.empty()
        st.info(f"🏁 **Hoàn thành!** Đã xử lý thành công **{success_count}/{len(uploaded_files)}** tài liệu.")
        # Xóa cache để danh sách được cập nhật
        st.cache_data.clear()

st.divider()

# Danh sách tài liệu hiện có
st.header("📖 Tài liệu hiện có trong hệ thống")

documents = get_documents_cached()

if not documents:
    st.info("Hiện chưa có tài liệu nào trong cơ sở dữ liệu.")
else:
    # Bộ lọc
    all_topics = ["Tất cả"] + list(topic_options.keys())
    filter_topic = st.selectbox(
        "Lọc theo chủ đề:",
        options=all_topics,
        format_func=lambda x: "Tất cả" if x == "Tất cả" else topic_options.get(x, x)
    )

    # Lọc và hiển thị
    if filter_topic == "Tất cả":
        filtered_docs = documents
    else:
        filtered_docs = [doc for doc in documents if doc.get('topic') == filter_topic]
        
    st.write(f"Tìm thấy {len(filtered_docs)} tài liệu.")

    # Hiển thị 10 tài liệu mới nhất
    for doc in reversed(filtered_docs[-10:]):
        with st.container(border=True):
            st.markdown(f"**📄 {doc.get('name', 'N/A')}**")
            meta_info = (
                f"**Chủ đề:** {topic_options.get(doc.get('topic'), 'N/A')} | "
                f"**Kích thước:** {doc.get('file_size', 0) / 1024:.1f} KB | "
                f"**Số phần:** {doc.get('chunk_count', 0)}"
            )
            st.caption(meta_info)